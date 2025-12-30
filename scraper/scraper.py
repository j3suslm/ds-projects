import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
import re

def scraper():
    st.subheader('Wikipedia Scraper')
    
    search_term = st.text_input('', placeholder='Buscar en Wikipedia', autocomplete='off')

    if search_term:
        base_url = 'https://es.wikipedia.org/wiki/'
        url = base_url + search_term.replace(' ', '_')
        st.write(f'URL generada: {url}')

        try:
            response = requests.get(url)
            response.raise_for_status()
            st.write('**Respuesta del servidor Success 200**')
            soup = BeautifulSoup(response.text, 'html.parser')
            title = soup.find('h1').text
            st.write(f'**Título del artículo**: {title}')
            word_text = []
            paragraphs = soup.find_all('p')
            for paragraph in paragraphs[:10]:
                text = re.sub(r'\[\d+\]', '', paragraph.text)
                word_text.append(text)
                st.write(text)
            if st.button('Generar word', type='primary'):
                doc = Document()
                doc.add_heading(title, level=1)
                for p in word_text:
                    doc.add_paragraph(p)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label='Descargar documento',
                    icon=':material/download:',
                    data=buffer,
                    file_name='documento.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                )
        except requests.exceptions.RequestException as e:
            st.error(f'Error al realizar solicitud: {e}')
        except Exception as e:
            st.error(f'Error en la pagina: {e}')
